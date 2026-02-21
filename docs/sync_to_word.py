import os
import re
from docx import Document
from datetime import datetime

# Installation pre-requisite: pip install python-docx

SOURCE_FILE = "SESSION_SUMMARY.md"
TARGET_DIR = "docs/sessions"
DOC_NAME = f"Session_Summary_{datetime.now().strftime('%Y_%m_%d')}.docx"

def convert_md_to_docx():
    if not os.path.exists(TARGET_DIR):
        os.makedirs(TARGET_DIR)
    
    if not os.path.exists(SOURCE_FILE):
        print(f"Error: {SOURCE_FILE} not found.")
        return

    doc = Document()
    doc.add_heading('Project Session Summaries', 0)

    with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
        content = f.readlines()

    for line in content:
        line = line.strip()
        if not line:
            continue
            
        # Very basic markdown parsing for headers and lists
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and '**:' in line:
            # Bold key: value
            p = doc.add_paragraph()
            parts = line.split('**:')
            run = p.add_run(parts[0].replace('**', '') + ':')
            run.bold = True
            p.add_run(parts[1])
        else:
            doc.add_paragraph(line)

    target_path = os.path.join(TARGET_DIR, DOC_NAME)
    doc.save(target_path)
    print(f"Successfully exported {SOURCE_FILE} to {target_path}")
    print("You can now safely drag and drop this file into WhatsApp or Teams!")

if __name__ == "__main__":
    convert_md_to_docx()
