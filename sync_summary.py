#!/usr/bin/env python3
"""
sync_summary.py - Convertit SESSION_SUMMARY.md en .docx et envoie sur WhatsApp

Ce script automatise:
1. La conversion du resume de session en document Word
2. L'envoi du resume vers WhatsApp pour l'equipe
3. Le stockage dans EMPIRE_DOCS pour synchronisation cloud

Usage:
    python sync_summary.py [--whatsapp] [--project PROJECT_NAME]

Configuration WhatsApp (fichier .env ou variables d'environnement):
    WHATSAPP_API_URL: URL de l'API WhatsApp (CallMeBot, Twilio, ou autre)
    WHATSAPP_PHONE: Numero de telephone destinataire
    WHATSAPP_APIKEY: Cle API (si necessaire)

Dependances:
    pip install python-docx requests python-dotenv
"""

import argparse
import os
import re
import json
import requests
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ python-docx non installe. Run: pip install python-docx")
    exit(1)

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass


# Configuration par defaut
EMPIRE_DOCS_DIR = Path(__file__).parent / "EMPIRE_DOCS"
PROJECT_NAME = os.getenv("PROJECT_NAME", "Charmed")


def parse_session_summary(content: str) -> dict:
    """Parse le contenu du SESSION_SUMMARY.md en structure dict"""
    sessions = []
    
    # Split par date (format: # Session Summary — YYYY-MM-DD)
    session_pattern = r'# Session Summary — (\d{4}-\d{2}-\d{2})'
    parts = re.split(session_pattern, content)
    
    # parts[0] est le header, puis alterne date/contenu
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            date = parts[i]
            body = parts[i + 1].strip()
            
            session = {
                'date': date,
                'editor': extract_field(body, r'\*\*Editor\*\*:\s*(.+)'),
                'fr': extract_section(body, 'Français'),
                'en': extract_section(body, 'English'),
                'tests': extract_field(body, r'\*\*Tests\*\*:\s*(.+)'),
                'blockers': extract_field(body, r'\*\*Blockers\*\*:\s*(.+)'),
            }
            sessions.append(session)
    
    return {'sessions': sessions}


def extract_field(text: str, pattern: str) -> str:
    """Extrait un champ specifique du texte"""
    match = re.search(pattern, text, re.DOTALL)
    return match.group(1).strip() if match else ''


def extract_section(text: str, header: str) -> dict:
    """Extrait une section (FR ou EN) du texte"""
    # Trouver le debut de la section (avec ou sans emoji)
    start_pattern = rf'## [🇫🇷🇬🇧]*\s*{header}'
    start_match = re.search(start_pattern, text)
    
    if not start_match:
        return {}
    
    start_pos = start_match.end()
    
    # Trouver la fin (prochain ## ou ---)
    end_match = re.search(r'\n---|\n## ', text[start_pos:])
    end_pos = start_pos + end_match.start() if end_match else len(text)
    
    section_text = text[start_pos:end_pos].strip()
    
    return {
        'what_done': extract_list_field(section_text, r'\*\*(?:Ce qui a été fait|What was done)\s*[:：]?\s*\n(.+?)(?=\n\*\*|\n\n\*\*|\Z)'),
        'files_changed': extract_list_field(section_text, r'\*\*(?:Fichiers modifiés|Files changed)\s*[:：]?\s*\n(.+?)(?=\n\*\*|\n\n\*\*|\Z)'),
        'next_steps': extract_list_field(section_text, r'\*\*(?:Étapes suivantes|Next steps)\s*[:：]?\s*\n(.+?)(?=\n\*\*|\n\n\*\*|\Z)'),
        'initiatives': extract_list_field(section_text, r'\*\*(?:Initiatives données|Initiatives given)\s*[:：]?\s*\n(.+?)(?=\n\*\*|\n\n\*\*|\Z)'),
    }


def extract_list_field(text: str, pattern: str) -> list:
    """Extrait une liste d'elements d'un champ"""
    match = re.search(pattern, text, re.DOTALL)
    if not match:
        return []
    
    content = match.group(1).strip()
    # Extraire les elements de liste
    items = re.findall(r'[-*]\s*(.+?)(?=\n[-*]|\n\n|\Z)', content, re.DOTALL)
    return [item.strip() for item in items if item.strip()]


def create_document(parsed_data: dict, output_path: Path, project_name: str):
    """Cree un document Word à partir des donnees parsees"""
    doc = Document()
    
    # Titre principal
    title = doc.add_heading(f'{project_name} - Session Summaries', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date de generation
    gen_date = doc.add_paragraph(f'Genere le: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    gen_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for session in parsed_data['sessions']:
        # Separateur
        doc.add_paragraph('─' * 50)
        
        # En-tete session
        doc.add_heading(f'Session — {session["date"]}', level=1)
        
        if session['editor']:
            p = doc.add_paragraph()
            p.add_run('Editor: ').bold = True
            p.add_run(session['editor'])
        
        # Section Francaise
        if session['fr'] and (session['fr'].get('what_done') or session['fr'].get('files_changed')):
            doc.add_heading('Francais', level=2)
            add_session_content(doc, session['fr'])
        
        # Section Anglaise
        if session['en'] and (session['en'].get('what_done') or session['en'].get('files_changed')):
            doc.add_heading('English', level=2)
            add_session_content(doc, session['en'])
        
        # Tests et Blockers
        if session['tests']:
            p = doc.add_paragraph()
            p.add_run('Tests: ').bold = True
            p.add_run(session['tests'])
        
        if session['blockers']:
            p = doc.add_paragraph()
            p.add_run('Blockers: ').bold = True
            p.add_run(session['blockers'])
    
    # Sauvegarder
    doc.save(output_path)
    print(f"Document cree: {output_path}")
    return output_path


def add_session_content(doc, section: dict):
    """Ajoute le contenu d'une section au document"""
    if section.get('what_done'):
        p = doc.add_paragraph()
        p.add_run('Ce qui a ete fait:').bold = True
        for item in section['what_done']:
            doc.add_paragraph(item, style='List Bullet')
    
    if section.get('files_changed'):
        p = doc.add_paragraph()
        p.add_run('Fichiers modifies:').bold = True
        for item in section['files_changed']:
            doc.add_paragraph(item, style='List Bullet')
    
    if section.get('next_steps'):
        p = doc.add_paragraph()
        p.add_run('Etapes suivantes:').bold = True
        for item in section['next_steps']:
            doc.add_paragraph(item, style='List Bullet')


def format_whatsapp_message(session: dict, project_name: str) -> str:
    """Formate le message pour WhatsApp"""
    lines = [
        f"📊 *{project_name}* - Session {session['date']}",
        f"📝 Editor: {session.get('editor', 'N/A')}",
        "",
        "✅ *Ce qui a ete fait:*",
    ]
    
    if session.get('fr', {}).get('what_done'):
        for item in session['fr']['what_done'][:5]:  # Max 5 items
            lines.append(f"• {item[:100]}")  # Max 100 chars par item
    
    lines.append("")
    lines.append("📁 *Fichiers modifies:*")
    if session.get('fr', {}).get('files_changed'):
        files = session['fr']['files_changed'][:5]
        lines.append(", ".join(files))
    
    lines.append("")
    lines.append("⏭️ *Prochaines etapes:*")
    if session.get('fr', {}).get('next_steps'):
        for item in session['fr']['next_steps'][:3]:
            lines.append(f"• {item[:100]}")
    
    if session.get('blockers'):
        lines.append(f"")
        lines.append(f"⚠️ *Blockers:* {session['blockers']}")
    
    lines.append("")
    lines.append(f"🤖 Genere automatiquement")
    
    return "\n".join(lines)


def send_to_whatsapp(message: str, phone: str, api_url: str, apikey: Optional[str] = None) -> bool:
    """
    Envoie le message sur WhatsApp
    
    Supporte plusieurs APIs:
    - CallMeBot: https://api.callmebot.com/whatsapp.php
    - Twilio: https://api.twilio.com/2010-04-01/Accounts/{AccountSid}/Messages
    - Custom webhook
    """
    try:
        # CallMeBot format
        if "callmebot" in api_url.lower():
            params = {
                "phone": phone,
                "text": message,
                "apikey": apikey,
            }
            response = requests.get(api_url, params=params, timeout=30)
        
        # Twilio format
        elif "twilio" in api_url.lower():
            headers = {"Content-Type": "application/x-www-form-urlencoded"}
            data = {
                "To": f"whatsapp:{phone}",
                "From": f"whatsapp:{os.getenv('TWILIO_PHONE')}",
                "Body": message,
            }
            response = requests.post(
                api_url,
                auth=(os.getenv("TWILIO_SID"), os.getenv("TWILIO_TOKEN")),
                data=data,
                headers=headers,
                timeout=30
            )
        
        # Custom webhook (format generique)
        else:
            headers = {"Content-Type": "application/json"}
            data = {
                "phone": phone,
                "message": message,
                "project": PROJECT_NAME,
                "timestamp": datetime.now().isoformat(),
            }
            if apikey:
                headers["Authorization"] = f"Bearer {apikey}"
            response = requests.post(api_url, json=data, headers=headers, timeout=30)
        
        if response.status_code == 200:
            print(f"✅ Message WhatsApp envoye a {phone}")
            return True
        else:
            print(f"❌ Erreur WhatsApp: {response.status_code} - {response.text}")
            return False
            
    except requests.RequestException as e:
        print(f"❌ Erreur reseau WhatsApp: {e}")
        return False


def get_latest_session(parsed_data: dict) -> Optional[dict]:
    """Recupere la session la plus recente"""
    if not parsed_data['sessions']:
        return None
    return parsed_data['sessions'][0]


def main():
    parser = argparse.ArgumentParser(description='Sync SESSION_SUMMARY.md to .docx et WhatsApp')
    parser.add_argument('--whatsapp', '-w', action='store_true', help='Envoyer sur WhatsApp')
    parser.add_argument('--project', '-p', default=PROJECT_NAME, help='Nom du projet')
    parser.add_argument('--output', '-o', default=str(EMPIRE_DOCS_DIR), help='Dossier de sortie')
    parser.add_argument('--phone', default=os.getenv('WHATSAPP_PHONE'), help='Numero WhatsApp')
    args = parser.parse_args()
    
    # Lire SESSION_SUMMARY.md
    summary_path = Path(__file__).parent / 'SESSION_SUMMARY.md'
    
    if not summary_path.exists():
        print(f"❌ Fichier non trouve: {summary_path}")
        return 1
    
    print(f"📖 Lecture de {summary_path}")
    content = summary_path.read_text(encoding='utf-8')
    
    # Parser
    print("🔄 Parsing...")
    parsed_data = parse_session_summary(content)
    print(f"   → {len(parsed_data['sessions'])} session(s) trouvee(s)")
    
    # Creer le dossier de sortie
    output_dir = Path(args.output)
    output_dir.mkdir(exist_ok=True)
    
    # Generer le nom du fichier avec la date du jour
    today = datetime.now().strftime("%Y%m%d")
    output_path = output_dir / f'{args.project}_Session_{today}.docx'
    
    # Creer le document
    create_document(parsed_data, output_path, args.project)
    
    # Envoyer sur WhatsApp si demande
    if args.whatsapp:
        api_url = os.getenv('WHATSAPP_API_URL')
        apikey = os.getenv('WHATSAPP_APIKEY')
        phone = args.phone
        
        if not all([api_url, phone]):
            print("❌ Configuration WhatsApp manquante. Voir .env:")
            print("   WHATSAPP_API_URL=https://api.callmebot.com/whatsapp.php")
            print("   WHATSAPP_PHONE=+336XXXXXXXX")
            print("   WHATSAPP_APIKEY=your_api_key")
        else:
            latest = get_latest_session(parsed_data)
            if latest:
                message = format_whatsapp_message(latest, args.project)
                send_to_whatsapp(message, phone, api_url, apikey)
    
    print(f"\n✅ Synchronisation terminee!")
    print(f"📄 Document: {output_path}")
    return 0


if __name__ == '__main__':
    exit(main())